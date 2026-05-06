from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from pydantic import BaseModel
from sqlalchemy.orm import Session
from typing import Optional
from db import get_db
from db.models import StockItem, ProductPackage
from api.auth import get_current_admin

router = APIRouter(prefix="/stock", tags=["stock"])


class StockBulkAdd(BaseModel):
    package_id: int
    items: str  # newline-separated data


class StockAdd(BaseModel):
    package_id: int
    data: str


@router.get("/package/{pkg_id}", dependencies=[Depends(get_current_admin)])
def list_stock(pkg_id: int, sold: Optional[bool] = None, db: Session = Depends(get_db)):
    q = db.query(StockItem).filter(StockItem.package_id == pkg_id)
    if sold is not None:
        q = q.filter(StockItem.is_sold == sold)
    items = q.order_by(StockItem.created_at.desc()).all()
    return [
        {
            "id": i.id,
            "data": i.data,
            "is_sold": i.is_sold,
            "sold_at": i.sold_at.isoformat() if i.sold_at else None,
            "order_id": i.order_id,
            "created_at": i.created_at.isoformat() if i.created_at else None,
        }
        for i in items
    ]


@router.post("/bulk", dependencies=[Depends(get_current_admin)])
def bulk_add_stock(data: StockBulkAdd, db: Session = Depends(get_db)):
    pkg = db.query(ProductPackage).filter(ProductPackage.id == data.package_id).first()
    if not pkg:
        raise HTTPException(status_code=404, detail="Package not found")

    lines = [line.strip() for line in data.items.strip().splitlines() if line.strip()]
    items = [StockItem(package_id=data.package_id, data=line) for line in lines]
    db.bulk_save_objects(items)
    db.commit()
    return {"added": len(items)}


@router.post("/", dependencies=[Depends(get_current_admin)])
def add_single_stock(data: StockAdd, db: Session = Depends(get_db)):
    pkg = db.query(ProductPackage).filter(ProductPackage.id == data.package_id).first()
    if not pkg:
        raise HTTPException(status_code=404, detail="Package not found")
    item = StockItem(package_id=data.package_id, data=data.data)
    db.add(item)
    db.commit()
    db.refresh(item)
    return {"id": item.id, "data": item.data}


@router.delete("/{item_id}", dependencies=[Depends(get_current_admin)])
def delete_stock(item_id: int, db: Session = Depends(get_db)):
    item = db.query(StockItem).filter(StockItem.id == item_id, StockItem.is_sold == False).first()
    if not item:
        raise HTTPException(status_code=404, detail="Stock item not found or already sold")
    db.delete(item)
    db.commit()
    return {"ok": True}


@router.get("/count/{pkg_id}")
def stock_count(pkg_id: int, db: Session = Depends(get_db)):
    count = db.query(StockItem).filter(
        StockItem.package_id == pkg_id,
        StockItem.is_sold == False
    ).count()
    return {"package_id": pkg_id, "available": count}


@router.post("/upload/{package_id}", dependencies=[Depends(get_current_admin)])
async def upload_stock_file(package_id: int, file: UploadFile = File(...), db: Session = Depends(get_db)):
    pkg = db.query(ProductPackage).filter(ProductPackage.id == package_id).first()
    if not pkg:
        raise HTTPException(status_code=404, detail="Package not found")

    filename = file.filename or ""
    content = await file.read()
    lines = []

    if filename.lower().endswith(".docx"):
        try:
            from docx import Document
            from io import BytesIO
            doc = Document(BytesIO(content))
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # A paragraph may contain multiple lines
                    for line in text.splitlines():
                        line = line.strip()
                        if line:
                            lines.append(line)
            # Also check tables in the docx
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for line in cell.text.strip().splitlines():
                            line = line.strip()
                            if line:
                                lines.append(line)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Lỗi đọc file .docx: {str(e)}")
    elif filename.lower().endswith(".txt") or filename.lower().endswith(".csv"):
        text = content.decode("utf-8", errors="replace")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
    else:
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .txt, .csv, .docx")

    if not lines:
        raise HTTPException(status_code=400, detail="File không có dữ liệu hợp lệ")

    items = [StockItem(package_id=package_id, data=line) for line in lines]
    db.bulk_save_objects(items)
    db.commit()
    return {"added": len(items)}
